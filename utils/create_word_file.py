import io

from docx import Document
from datetime import datetime


def create_word_file(name, phone, data):
    doc = Document()
    doc.add_heading('Смета на паркетные работы', 1)
    doc.add_heading('Предварительная смета, счетом не является', 9)

    doc.add_paragraph(f'Заказчик: {name}, {phone}')
    doc.add_paragraph(f'Площадь помещения: {data["area"]} м²')

    total = 0

    for category, items in data['categories'].items():
        doc.add_heading(category, level=2)
        for item in items:
            item_name = item['name']
            price = int(item['price'])
            if item['type'] in ['checkbox', 'radio']:
                cost = price * int(data['area'])
                total += cost
                doc.add_paragraph(f'{item_name} — {cost} ₽')
            elif item['type'] == 'number':
                value = item.get('value', 0)
                cost = price * value
                total += cost
                doc.add_paragraph(f'{item_name} {value} × {price} ₽ = {cost} ₽')

    doc.add_paragraph(f'Итого: {total} ₽', style='Intense Quote')

    today = datetime.today().strftime('%d.%m.%Y')
    doc.add_paragraph(f'\nПодготовка сметы: Николаенко Р.М.\nДата: {today}')

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
